from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_resume():
    # Create a new Document
    doc = Document()

    # Set up the document
    doc.add_heading('Anil Kumar B', 0)

    # Contact Information
    p = doc.add_paragraph()
    p.add_run('571-508-9800 | anilraobalguri1998@gmail.com | GitHub: https://github.com/Anil4518').bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Experience
    doc.add_heading('Experience', level=1)

    # Western Illinois University
    p = doc.add_paragraph()
    p.add_run('Western Illinois University, Macomb, IL\t').bold = True
    p.add_run('Jan 2022 - Aug 2023\n')
    p.add_run('Graduate Assistant\n')
    p.add_run('• Collected and managed large datasets for academic and research projects, ensuring data integrity.\n')
    p.add_run('• Conducted statistical analyses using Python and R to support research initiatives.\n')
    p.add_run('• Generated comprehensive reports and visualizations to present findings to faculty.\n')
    p.add_run('• Assisted in database management and utilized SQL for efficient data retrieval.')

    # Amazon
    p = doc.add_paragraph()
    p.add_run('Amazon, Hyderabad TS\t').bold = True
    p.add_run('June 2020 - Dec 2021\n')
    p.add_run('Data Analyst\n')
    p.add_run('• Led data-driven projects and automated data processing using Python and SQL, improving efficiency by 30%.\n')
    p.add_run('• Managed data storage with Amazon Redshift, MySQL, and S3, ensuring data integrity and reducing retrieval time by 20%.\n')
    p.add_run('• Utilized AWS Glue and Apache Airflow for ETL processes, increasing ETL pipeline reliability by 25%.\n')
    p.add_run('• Enhanced decision-making with predictive analytics, boosting customer segmentation and targeted marketing strategies by 15%.')

    # Relevant Projects
    doc.add_heading('Relevant Projects', level=1)

    p = doc.add_paragraph()
    p.add_run('Stock Price Prediction\n').bold = True
    p.add_run('• Developed predictive model for stock prices using historical data and financial indicators.\n')
    p.add_run('• Collected and preprocessed data with Python and SQL.\n')
    p.add_run('• Trained models using Pandas, NumPy, scikit-learn, TensorFlow, and Keras.\n')
    p.add_run('• Created a web app for predictions using JavaScript and React.js.\n')
    p.add_run('• Presented findings with interactive visualizations in Tableau.')

    # Skills
    doc.add_heading('Skills', level=1)

    p = doc.add_paragraph()
    p.add_run('• Data Analysis: SQL, MS Excel, Python (Pandas, NumPy, Matplotlib), Alteryx, SAS, Tableau\n')
    p.add_run('• Data Modeling: Data warehousing, customized solutions for data management\n')
    p.add_run('• Machine Learning: scikit-learn, TensorFlow, Keras, Flask, Django\n')
    p.add_run('• Programming: Python, SQL, JavaScript\n')
    p.add_run('• Methodologies: Agile')

    # Education
    doc.add_heading('Education', level=1)

    p = doc.add_paragraph()
    p.add_run('Master’s Degree\n').bold = True
    p.add_run('Applied Statistics and Data Analytics, Western Illinois University, IL (Jan 2022 - Dec 2023)\n')

    p = doc.add_paragraph()
    p.add_run('Bachelor’s Degree\n').bold = True
    p.add_run('Mechanical Engineering, University College of Engineering, Osmania University, Hyderabad (Aug 2016 - Sept 2020)')

    # Save the document
    doc.save('./ATS_Friendly_Resume.docx')

if __name__ == "__main__":
    create_resume()
